"""Generate Bikram Parida's updated CV as a .docx file."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


OUTPUT = Path(__file__).resolve().parent.parent / "Bikram_Parida_CV.docx"


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    add_bottom_border(p)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.size = Pt(10.5)


def add_body(doc, text: str, bold: bool = False, italic: bool = False, space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)


def build() -> Path:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Header: Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    name_run = name.add_run("BIKRAM PARIDA")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Contact line
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)
    contact_run = contact.add_run("Mobile: +91-8249235053  |  Email: bikramparida125@gmail.com")
    contact_run.font.size = Pt(10.5)
    contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    add_bottom_border(contact)

    # Career Objective
    add_heading(doc, "Career Objective")
    add_body(
        doc,
        "To secure a challenging position in a reputable organization where I can expand my "
        "learning, knowledge, and skills. Seeking a responsible career opportunity that allows "
        "me to fully utilize my training and abilities while making a significant contribution "
        "to the success of the company.",
    )

    # Work Experience
    add_heading(doc, "Work Experience")

    # Job 1 - current
    role1 = doc.add_paragraph()
    role1.paragraph_format.space_after = Pt(0)
    r = role1.add_run("Desktop Engineer  ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = role1.add_run("| Sysnet Global Technologies (deployed at Volvo)")
    r2.font.size = Pt(11)
    add_body(doc, "June 2025 - Present", italic=True, space_after=2)
    add_bullet(doc, "Provide on-site desktop support and IT services for Volvo end users.")
    add_bullet(doc, "Handle hardware and software issues, system configuration, and user-account management.")

    # Job 2
    role2 = doc.add_paragraph()
    role2.paragraph_format.space_before = Pt(6)
    role2.paragraph_format.space_after = Pt(0)
    r = role2.add_run("Desktop Support Engineer  ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = role2.add_run("| Moderncrow Technologies Pvt. Ltd.")
    r2.font.size = Pt(11)
    add_body(doc, "March 2023 - March 2025", italic=True, space_after=2)
    add_bullet(doc, "Delivered desktop support services including OS installation, troubleshooting, and hardware repair.")
    add_bullet(doc, "Managed remote desktop assistance, ticket resolution, and end-user account setup.")

    # Job Proficiency Skills
    add_heading(doc, "Job Proficiency Skills")
    skills = [
        "OS installation including Windows 10, 11, and Windows Server.",
        "Application software installation and activation.",
        "Desktop rework and testing.",
        "Fault diagnosis, repair, and servicing.",
        "Hands-on experience repairing HP, Dell, Microsoft, Lenovo, and Intel systems.",
        "Desktop assembly and disassembly.",
        "File and printer sharing.",
        "Remote desktop support.",
        "Applying BitLocker on specific drives.",
        "Configuring and troubleshooting MS Outlook.",
        "Resolving issues through ticketing tools.",
        "Creating local users on client machines.",
        "Creating local and domain users.",
        "Configuring IP addresses.",
        "Hands-on experience installing and configuring Windows Server 2012 and 2016.",
    ]
    for s in skills:
        add_bullet(doc, s)

    # Education
    add_heading(doc, "Educational Qualifications")
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    hdr = table.rows[0].cells
    headers = ["Qualification", "School / College", "Year"]
    for cell, text in zip(hdr, headers):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3A5F")

    rows = [
        ("Matriculation", "R.N. Academy", "2019"),
        ("12th (Higher Secondary)", "Jogendra Higher Secondary School", "2021"),
        ("Graduation", "Utkal University", "2025"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, val in zip(cells, row_data):
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10.5)

    # Technical Skills
    add_heading(doc, "Technical Skills")

    def add_subheading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)

    add_subheading("Microsoft Office")
    add_bullet(doc, "Microsoft Word, Excel, PowerPoint, and Outlook.")

    add_subheading("Operating Systems")
    for s in [
        "Installation of Windows 10 and 11.",
        "Windows upgrading, clean install, and dual booting.",
        "Software and driver installation; user-account management.",
        "Disk management, network settings, file sharing, and troubleshooting.",
        "Remote assistance, BitLocker, Windows Update, backup and restore.",
        "Data recovery, printer installation, sharing, and troubleshooting.",
        "Windows password reset and BIOS setup.",
    ]:
        add_bullet(doc, s)

    add_subheading("Hardware")
    for s in [
        "Desktop assembly, disassembly, and troubleshooting.",
        "Identification, installation, and troubleshooting of processors, hard disks, RAM, and SMPS units.",
    ]:
        add_bullet(doc, s)

    add_subheading("Networking")
    for s in [
        "IP configuration.",
        "OSI model, TCP/IP model, and structured cabling.",
        "Configuration of DNS, DHCP, WDS, and FTP.",
    ]:
        add_bullet(doc, s)

    # Declaration
    add_heading(doc, "Declaration")
    add_body(
        doc,
        "I hereby declare that all the information stated above is true and correct to the best "
        "of my knowledge and belief.",
        space_after=10,
    )

    # Signature line
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(12)
    tab_stops = sig.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    sig_run = sig.add_run("Place: ____________\tBikram Parida")
    sig_run.font.size = Pt(10.5)

    date_p = doc.add_paragraph()
    date_run = date_p.add_run("Date:  ____________")
    date_run.font.size = Pt(10.5)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Saved: {path}")
